"""Regenerate every file in this folder.

The `.txt` files are readable in any editor; the `.xlsx`, `.docx` and `.pdf`
files are binary, so this script is the readable source for them — edit here
and re-run rather than hand-editing a binary.

    cd backend
    uv run python ../sample/build_samples.py

WHAT IS REAL AND WHAT IS NOT
----------------------------
The **questionnaire** is modelled on real published standards. Its questions
are worded from the Capital Markets Malaysia *Simplified ESG Disclosure Guide
(SEDG) Version 2* (July 2025, 38 disclosures across 15 topics), cross-checked
against the nine common sustainability matters Bursa Malaysia requires of Main
Market issuers. Disclosure codes such as `SEDG-E1.1` are the real ones. This
matters: a demo questionnaire invented from scratch teaches you nothing about
whether the product can answer the questions customers actually send.

The **evidence is entirely synthetic** and describes a company that does not
exist (AGENTS.md §3.1 — synthetic data only, no exceptions). Two files marked
`reference-` carry real published figures with their source named inside the
file, and nothing else does.

FILE FORMATS AND WHY EACH IS USED
---------------------------------
The backend has one parser per format and they chunk differently, which changes
what a source citation looks like:

    .pdf   one chunk per page           -> location is a page number
    .docx  one chunk per heading         -> location is a heading path
    .xlsx  one chunk per row             -> location is a sheet + cell range
    .txt   one chunk per non-blank line  -> location is a line

So each file is laid out one fact per chunk. A one-fact-per-line text file
produces a precise citation; a wall of prose produces a vague one.

THE THREE TIERS
---------------
Evidence filenames are prefixed by how trustworthy the document is, because the
whole point of the sample set is that a reviewer has to tell them apart:

    A-*   sound. Internally consistent, inside the reporting period, scope
          stated, and it answers the question that cites it. A reviewer should
          be able to confirm the answer from it.
    B-*   uncertain. Real documents with a real gap — half the year, no unit,
          a group figure where a site figure was asked for, an admitted
          estimate. Not wrong, not sufficient.
    C-*   wrong. Contradicts a sound document, or is three years stale, or
          belongs to a different legal entity, or cannot be read at all.

WHAT THE RULE ENGINE CAN AND CANNOT SEE (read this before demoing)
------------------------------------------------------------------
The engine grades A, B and C almost identically — nearly everything comes out
`PARTIAL`. That is not a flaw in this sample set; it is the current ceiling of
the pipeline, and these files are built to make it visible:

* `evidence_links.value` is never populated. The matcher writes only
  `chunk_id`, `claim_supported` and `quoted_excerpt`
  (`packages/ai-pipeline/src/ai_pipeline/analyze.py`). `CONFLICTING` requires
  two links with the same scope and period and *different values*, so the
  contradiction between `A-03` and `C-01` — 12.6 t against 18.4 t of scheduled
  waste for the same year — is invisible to the engine. A human has to catch it.
* `documents.source_date` cannot be set on upload (the endpoint takes only
  `document_type`), so `OUTDATED` never fires either. `C-02` is three years
  stale and the engine will not say so.
* No endpoint moves an `evidence_links` row to `ACCEPTED`, so every link stays
  `CANDIDATE` and `_partial_reasons` always returns at least
  `REASON_NOT_ACCEPTED`. `VERIFIED` is therefore unreachable through the API.

Reachable today: `MISSING`, `PARTIAL`, and `NOT_APPLICABLE` (a human action).
Everything in the C tier is a gap the *reviewer* closes, not the engine. Two
questions are also worded so that nothing matches at all, which is what
`MISSING` looks like — a demo where every row lights up green teaches nothing.
"""

from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parent
QUESTIONNAIRE_DIR = ROOT / "questionnaire"
EVIDENCE_DIR = ROOT / "evidence"

COMPANY = "Tenggara Precision Sdn. Bhd."
SITE = "Lot 88, Jalan Perusahaan 4, Bukit Raja Industrial Estate, 41050 Klang, Selangor"
PERIOD = "FY2025 (1 January 2025 to 31 December 2025)"

# The figures every sound document agrees on. Kept in one place so the A tier
# reconciles by construction and the C tier can disagree on purpose.
ELECTRICITY_KWH = 1_847_300
GRID_FACTOR_2024 = 0.740  # tCO2e/MWh, Peninsular. Real; see the reference file.
SCOPE_2_TCO2E = round(ELECTRICITY_KWH / 1000 * GRID_FACTOR_2024, 1)  # 1367.0
SCOPE_1_TCO2E = 58.8
WASTE_TOTAL_T = 214.7
WASTE_DIVERTED_T = 138.2
WASTE_DISPOSED_T = 76.5
SCHEDULED_WASTE_T = 12.6
HEADCOUNT = 268
TRAINING_HOURS_TOTAL = 3216


# --------------------------------------------------------------------------- #
# Questionnaire
# --------------------------------------------------------------------------- #

# (external_question_id, question_text, section, is_required)
#
# Wording follows SEDG v2 disclosure text. It is also chosen so the keyword
# mapper in packages/ai-pipeline/src/ai_pipeline/sedg_taxonomy.py can reach a
# topic — that taxonomy is a documented placeholder, not the real SEDG codes,
# so the overlap has to come from the words themselves ("Scope 1", "kWh",
# "water withdrawal", "LTIFR", "anti-corruption").
QUESTIONS: list[tuple[str, str, str, bool]] = [
    # -- Environmental ------------------------------------------------------
    (
        "Q-E-01",
        "SEDG-E1.1: Report total Scope 1 (direct) GHG emissions in metric tonnes of "
        "CO2 equivalent for the reporting period.",
        "Environmental",
        True,
    ),
    (
        "Q-E-02",
        "SEDG-E1.2: Report total Scope 2 (indirect) GHG emissions in metric tonnes of "
        "CO2 equivalent, stating the grid emission factor used and its source.",
        "Environmental",
        True,
    ),
    (
        "Q-E-03",
        "SEDG-E1.7: Report total Scope 1 and 2 GHG emissions intensity in metric tonnes "
        "of CO2 equivalent per unit of production.",
        "Environmental",
        False,
    ),
    (
        "Q-E-04",
        "SEDG-E2.1: Report total energy consumption in kWh, broken down into electricity "
        "purchased from the grid, renewable fuel sources and non-renewable fuel sources.",
        "Environmental",
        True,
    ),
    (
        "Q-E-05",
        "SEDG-E2.2: Report the reduction in energy consumption achieved as a direct "
        "result of conservation and efficiency initiatives, in kWh.",
        "Environmental",
        False,
    ),
    (
        "Q-E-06",
        "SEDG-E3.1: Report the total water withdrawal from all areas for the reporting "
        "period, with a breakdown by source.",
        "Environmental",
        True,
    ),
    (
        "Q-E-07",
        "SEDG-E4.1: Report total waste generated, total waste diverted from disposal and "
        "total waste directed to disposal, in metric tonnes.",
        "Environmental",
        True,
    ),
    (
        "Q-E-08",
        "SEDG-E4.2: Report total hazardous waste and scheduled waste generated in metric "
        "tonnes, identified by waste code.",
        "Environmental",
        True,
    ),
    (
        "Q-E-09",
        "SEDG-E5.2: Report the percentage of recycled input materials used to manufacture "
        "the company's primary products.",
        "Environmental",
        False,
    ),
    # -- Social -------------------------------------------------------------
    (
        "Q-S-01",
        "SEDG-S1.1: Report the number and nature of child labour and forced labour "
        "incidents recorded during the reporting period, if any.",
        "Social",
        True,
    ),
    (
        "Q-S-02",
        "SEDG-S2.1: Report the average hours of training per employee for the reporting "
        "period.",
        "Social",
        True,
    ),
    (
        "Q-S-03",
        "SEDG-S2.2: Report the total number of employees and the employee turnover rate "
        "for the reporting period.",
        "Social",
        True,
    ),
    (
        "Q-S-04",
        "SEDG-S3.1: Report the percentage of employees by gender and by age band, "
        "including gender diversity in management positions.",
        "Social",
        True,
    ),
    (
        "Q-S-05",
        "SEDG-S4.1: Report the number of work-related fatalities and the number of "
        "work-related injuries, stating the LTIFR and the hours worked it is based on.",
        "Social",
        True,
    ),
    (
        "Q-S-06",
        "SEDG-S4.2: Report the total number and percentage of employees trained on "
        "health and safety standards during the reporting period.",
        "Social",
        False,
    ),
    # -- Governance ---------------------------------------------------------
    (
        "Q-G-01",
        "SEDG-G2.1: List the company's policy commitments, including the code of "
        "conduct, the anti-corruption policy and the whistleblowing policy, giving the "
        "approval date and the version in force.",
        "Governance",
        True,
    ),
    (
        "Q-G-02",
        "SEDG-G4.2: Report the total number and percentage of employees who received "
        "training on the company's anti-bribery and anti-corruption policy.",
        "Governance",
        True,
    ),
    (
        "Q-G-03",
        "SEDG-G5.1: Report the number and nature of substantiated complaints concerning "
        "breaches of customer privacy or loss of customer data, if any.",
        "Governance",
        True,
    ),
    # -- Deliberately unanswerable from this evidence set --------------------
    #
    # Nothing uploaded mentions either subject, and no taxonomy keyword
    # matches. These are what MISSING looks like, and a demo without any is a
    # demo that has quietly hidden the hard half of the job.
    (
        "Q-E-10",
        "SEDG-E5.1: List the materials and total weights used to package the company's "
        "primary products, in metric tonnes.",
        "Environmental",
        False,
    ),
    (
        "Q-S-07",
        "SEDG-S5.1: Report the total amount of community investments and donations made "
        "during the reporting period.",
        "Social",
        False,
    ),
]


def build_questionnaire() -> None:
    """The header row is load-bearing: the parser requires exactly
    `external_question_id | question_text | section | is_required` on row 1.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire"
    ws.append(["external_question_id", "question_text", "section", "is_required"])
    for qid, text, section, required in QUESTIONS:
        ws.append([qid, text, section, required])
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 100
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 12
    wb.save(QUESTIONNAIRE_DIR / "customer-esg-questionnaire-2026.xlsx")


def build_broken_questionnaire() -> None:
    """Deliberately wrong headers, to exercise the parse-failure path.

    Uploading this as a QUESTIONNAIRE must leave the document in
    NEEDS_MANUAL_REVIEW naming the missing headers, and must not create a
    single question. Silently producing zero questions would be the actual bug.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["No.", "Question", "Category", "Mandatory?"])
    ws.append(["1", "Report total annual electricity consumption.", "Environment", "Yes"])
    ws.append(["2", "Describe the environmental policy.", "Environment", "Yes"])
    wb.save(QUESTIONNAIRE_DIR / "broken-questionnaire-wrong-headers.xlsx")


# --------------------------------------------------------------------------- #
# A tier — sound evidence (PDF, one chunk per page)
# --------------------------------------------------------------------------- #

# Twelve months that sum to ELECTRICITY_KWH exactly. A reviewer who adds them
# up and gets a different number has found a bug in this generator, not a
# judgement call.
BILL_MONTHS = [
    ("January 2025", "01/01/2025 - 31/01/2025", 148_600),
    ("February 2025", "01/02/2025 - 28/02/2025", 139_200),
    ("March 2025", "01/03/2025 - 31/03/2025", 156_400),
    ("April 2025", "01/04/2025 - 30/04/2025", 151_900),
    ("May 2025", "01/05/2025 - 31/05/2025", 162_300),
    ("June 2025", "01/06/2025 - 30/06/2025", 158_700),
    ("July 2025", "01/07/2025 - 31/07/2025", 164_100),
    ("August 2025", "01/08/2025 - 31/08/2025", 159_800),
    ("September 2025", "01/09/2025 - 30/09/2025", 152_600),
    ("October 2025", "01/10/2025 - 31/10/2025", 157_400),
    ("November 2025", "01/11/2025 - 30/11/2025", 148_900),
    ("December 2025", "01/12/2025 - 31/12/2025", 147_400),
]

TARIFF_SEN_PER_KWH = 45.40  # Real published RP4 rate; see the reference file.


def build_electricity_bills_pdf() -> None:
    """Twelve monthly statements, one per page, so each month cites its own page.

    The tariff is the real published RP4 rate. The account number, premise and
    every kWh figure are invented.
    """
    assert sum(kwh for _, _, kwh in BILL_MONTHS) == ELECTRICITY_KWH, (
        "monthly bills must sum to the annual total the other A-tier documents use"
    )

    doc = pymupdf.open()
    for month, period, kwh in BILL_MONTHS:
        page = doc.new_page()
        amount = kwh * TARIFF_SEN_PER_KWH / 100
        lines = [
            "TENAGA NASIONAL BERHAD",
            "ELECTRICITY STATEMENT (SYNTHETIC SAMPLE - NOT A REAL BILL)",
            "",
            f"Account holder: {COMPANY}",
            f"Premise: {SITE}",
            "Account number: 21-0447-9931-08",
            f"Billing month: {month}",
            f"Billing period: {period}",
            "",
            f"Total consumption for the month: {kwh:,} kWh",
            f"Base tariff applied: {TARIFF_SEN_PER_KWH:.2f} sen/kWh",
            f"Amount charged: RM {amount:,.2f}",
            "",
            "This statement covers the Klang plant only. It does not include",
            "any other premise operated by the group.",
        ]
        page.insert_text((60, 70), "\n".join(lines), fontsize=11, lineheight=1.45)
    doc.save(EVIDENCE_DIR / "A-01-tnb-electricity-bills-fy2025.pdf")
    doc.close()


def build_unreadable_pdf() -> None:
    """A PDF with no extractable text, standing in for a photographed scan.

    Must land on NEEDS_MANUAL_REVIEW naming the reason, and must be excluded
    from the evidence-quality computation rather than silently counted.
    """
    doc = pymupdf.open()
    page = doc.new_page()
    # A drawn rectangle and nothing else: real page geometry, zero text.
    page.draw_rect(pymupdf.Rect(80, 80, 520, 700), color=(0.75, 0.75, 0.75), width=1.5)
    doc.save(EVIDENCE_DIR / "C-06-unreadable-scan-safety-records.pdf")
    doc.close()


# --------------------------------------------------------------------------- #
# A tier — sound evidence (XLSX, one chunk per row)
# --------------------------------------------------------------------------- #


def _write_xlsx(path: Path, sheet: str, header: list[str], rows: list[list], widths: list[int]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    ws.append(header)
    for row in rows:
        ws.append(row)
    for column, width in zip("ABCDEFGH", widths):
        ws.column_dimensions[column].width = width
    wb.save(path)


def build_ghg_inventory_xlsx() -> None:
    """Scope 1 and Scope 2 worked from the same electricity figure the bills show.

    The Scope 2 line is deliberately reproducible: 1,847,300 kWh is 1,847.3 MWh,
    times the 0.740 Peninsular factor is 1,367.0 tCO2e. A reviewer can redo it
    from `A-01` and `reference-malaysia-grid-emission-factor.txt` without
    trusting this sheet.
    """
    _write_xlsx(
        EVIDENCE_DIR / "A-02-ghg-inventory-fy2025.xlsx",
        "GHG Inventory FY2025",
        ["Scope", "Source", "Activity data", "Factor", "tCO2e", "Basis"],
        [
            [
                "Scope 1",
                "Standby diesel generator",
                "12,400 litres diesel",
                "2.68 kgCO2e/litre",
                33.2,
                "Fuel purchase records, Klang plant only",
            ],
            [
                "Scope 1",
                "LPG forklift fleet",
                "8,600 kg LPG",
                "2.98 kgCO2e/kg",
                25.6,
                "Cylinder delivery notes, Klang plant only",
            ],
            [
                "Scope 1 total",
                "Direct combustion controlled by the company",
                "",
                "",
                SCOPE_1_TCO2E,
                "Sum of the two rows above",
            ],
            [
                "Scope 2",
                "Purchased grid electricity",
                f"{ELECTRICITY_KWH:,} kWh = 1,847.3 MWh",
                f"{GRID_FACTOR_2024} tCO2e/MWh (Peninsular, 2024)",
                SCOPE_2_TCO2E,
                "Location-based. Ties to A-01 monthly statements",
            ],
            [
                "Factor note",
                "No published Peninsular grid factor exists for 2025 yet",
                "",
                "",
                "",
                "2024 provisional factor used; stated rather than assumed",
            ],
            [
                "Boundary",
                "Klang plant only",
                "",
                "",
                "",
                "Excludes Tenggara Logistics Sdn. Bhd. and the Ipoh site",
            ],
        ],
        [16, 42, 34, 34, 12, 46],
    )


def build_scheduled_waste_xlsx() -> None:
    """Scheduled-waste consignment summary. The A-tier waste figure.

    `C-01` contradicts the scheduled-waste total here for the same period. The
    engine cannot see that (it never extracts values), so the contradiction is
    left for a human — which is the point of shipping both.
    """
    rows = [
        ["SW 410", "Spent mineral oil from machining", 4.8, "2025", "Klang plant"],
        ["SW 305", "Spent lubricating oil", 2.9, "2025", "Klang plant"],
        ["SW 110", "Waste from electroplating rinse", 3.1, "2025", "Klang plant"],
        ["SW 409", "Contaminated containers and rags", 1.8, "2025", "Klang plant"],
        [
            "Total scheduled waste",
            "All codes above, FY2025",
            SCHEDULED_WASTE_T,
            "2025",
            "Klang plant",
        ],
        [
            "Non-hazardous waste",
            "General and packaging waste, FY2025",
            round(WASTE_TOTAL_T - SCHEDULED_WASTE_T, 1),
            "2025",
            "Klang plant",
        ],
        ["Total waste generated", "Hazardous plus non-hazardous", WASTE_TOTAL_T, "2025", "Klang plant"],
        ["Waste diverted from disposal", "Recycling and recovery", WASTE_DIVERTED_T, "2025", "Klang plant"],
        ["Waste directed to disposal", "Landfill and incineration", WASTE_DISPOSED_T, "2025", "Klang plant"],
    ]
    _write_xlsx(
        EVIDENCE_DIR / "A-03-scheduled-waste-consignment-fy2025.xlsx",
        "Waste FY2025",
        ["Waste code / line", "Description", "Metric tonnes", "Year", "Scope"],
        rows,
        [26, 44, 15, 10, 16],
    )


def build_training_register_xlsx() -> None:
    """Training hours per department, summing to a computable average.

    3,216 hours over 268 employees is exactly 12.0 hours each, so the answer to
    Q-S-02 can be checked rather than believed.
    """
    departments = [
        ("Production", 148, 1_628),
        ("Quality assurance", 42, 546),
        ("Maintenance", 31, 434),
        ("Warehouse and logistics", 26, 286),
        ("Administration and finance", 21, 322),
    ]
    rows = [[name, headcount, hours, round(hours / headcount, 1)] for name, headcount, hours in departments]
    total_head = sum(h for _, h, _ in departments)
    total_hours = sum(t for _, _, t in departments)
    assert total_head == HEADCOUNT and total_hours == TRAINING_HOURS_TOTAL, (
        "department rows must reconcile to the headline headcount and hours"
    )
    rows.append(["All departments, FY2025", total_head, total_hours, round(total_hours / total_head, 1)])
    _write_xlsx(
        EVIDENCE_DIR / "A-04-training-register-fy2025.xlsx",
        "Training FY2025",
        ["Department", "Employees", "Total training hours", "Average hours per employee"],
        rows,
        [32, 12, 22, 28],
    )


# --------------------------------------------------------------------------- #
# DOCX (one chunk per heading section)
# --------------------------------------------------------------------------- #


def _write_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    doc.save(path)


def build_anti_bribery_policy_docx() -> None:
    """A tier. Current, dated, version-stamped, and it names its own channel."""
    _write_docx(
        EVIDENCE_DIR / "A-05-anti-bribery-and-whistleblowing-policy-2025.docx",
        f"{COMPANY} — Anti-Bribery, Anti-Corruption and Whistleblowing Policy",
        [
            (
                "Version and approval",
                [
                    "Version 4.0, approved by the Board on 14 February 2025.",
                    "This version supersedes version 3.0 of 2022 and every earlier "
                    "statement on the same subject, including the anti-corruption "
                    "clauses in the 2019 Employee Handbook.",
                    "Next scheduled review: February 2027.",
                ],
            ),
            (
                "Scope",
                [
                    f"Applies to all employees, directors and contractors of {COMPANY} "
                    "at the Klang plant.",
                    "It does not cover Tenggara Logistics Sdn. Bhd., which maintains a "
                    "separate policy.",
                ],
            ),
            (
                "Anti-corruption commitment",
                [
                    "The company prohibits bribery, kickbacks and facilitation payments "
                    "in any form, whether offered directly or through a third party.",
                    "Gifts and hospitality above RM 500 require written approval from "
                    "the Managing Director before they are offered or accepted.",
                    "Adequate procedures are maintained in line with Section 17A of the "
                    "Malaysian Anti-Corruption Commission Act 2009.",
                ],
            ),
            (
                "Anti-corruption training",
                [
                    "All employees complete anti-bribery and anti-corruption training "
                    "within 30 days of joining and refresh it every two years.",
                    "In FY2025, 254 of 268 employees completed the training, which is "
                    "94.8 per cent. The remaining 14 joined in December 2025 and are "
                    "scheduled for the January 2026 intake.",
                ],
            ),
            (
                "Whistleblowing channel",
                [
                    "Reports may be made in confidence to whistleblow@tenggara-precision.example "
                    "or by calling the independent line on 1-800-88-0000.",
                    "The channel is operated by an external provider and reports may be "
                    "made anonymously.",
                    "Retaliation against a person who reports in good faith is a "
                    "disciplinary offence.",
                ],
            ),
            (
                "Confirmed incidents",
                [
                    "No confirmed incidents of corruption were recorded in FY2025.",
                    "One report was received in March 2025 concerning a procurement "
                    "gift; it was investigated and closed as unsubstantiated.",
                ],
            ),
        ],
    )


def build_superseded_handbook_docx() -> None:
    """C tier. A withdrawn policy that still reads like an authority.

    It contradicts A-05 on the whistleblowing channel and on the gift
    threshold. Nothing on its face says "withdrawn" except the date, which is
    exactly how stale policies cause wrong answers in real questionnaires.
    """
    _write_docx(
        EVIDENCE_DIR / "C-05-employee-handbook-2019-superseded.docx",
        f"{COMPANY} — Employee Handbook (2019 edition)",
        [
            (
                "Issue",
                [
                    "First edition, issued 3 June 2019.",
                    "Distributed to all employees on joining.",
                ],
            ),
            (
                "Reporting concerns",
                [
                    "Employees who wish to report misconduct should speak to their "
                    "immediate supervisor or to the Human Resources Manager in person.",
                    "There is no anonymous reporting channel.",
                ],
            ),
            (
                "Gifts",
                [
                    "Gifts from suppliers up to RM 1,000 may be accepted without "
                    "approval provided they are declared at the next department meeting.",
                ],
            ),
            (
                "Working hours and leave",
                [
                    "Standard working hours are 8.30am to 5.30pm, Monday to Friday.",
                    "Annual leave entitlement begins at 12 days and rises with service.",
                ],
            ),
        ],
    )


# --------------------------------------------------------------------------- #
# TXT (one chunk per non-blank line)
# --------------------------------------------------------------------------- #

TEXT_FILES: dict[str, list[str]] = {
    # ---- reference material: REAL published figures ----------------------
    "reference-malaysia-grid-emission-factor.txt": [
        "MALAYSIA GRID EMISSION FACTOR - REAL PUBLISHED REFERENCE DATA",
        "Figures below are real. Source: Grid Emission Factor (GEF) in Malaysia, published by the Energy Commission of Malaysia (Suruhanjaya Tenaga). Retrieved 2026-08-23. Content was rephrased for compliance with licensing restrictions.",
        "Unit: Gg CO2e/GWh, numerically equal to tCO2e/MWh.",
        "Provisional series 2022-2024: 2024 Peninsular 0.740, Sabah 0.539, Sarawak 0.199.",
        "Provisional series 2022-2024: 2023 Peninsular 0.760, Sabah 0.545, Sarawak 0.206.",
        "Provisional series 2022-2024: 2022 Peninsular 0.769, Sabah 0.531, Sarawak 0.199.",
        "Earlier series 2017-2022: 2022 Peninsular 0.774, Sabah 0.525, Sarawak 0.199.",
        "Earlier series 2017-2022: 2021 Peninsular 0.757, Sabah 0.524, Sarawak 0.198.",
        "Earlier series 2017-2022: 2020 Peninsular 0.821, Sabah 0.503, Sarawak 0.203.",
        "No Peninsular grid emission factor has been published for 2025 at the time of writing, so a 2025 Scope 2 calculation must state which earlier factor it used.",
        "DISCREPANCY WORTH NOTING: the two published series give different Peninsular values for 2022 - 0.769 provisional against 0.774 earlier. Both come from the same publisher. This is a real example of the problem this product exists to surface: two credible sources, one metric, two numbers, and no automatic way to decide which is right.",
        "A Scope 2 location-based calculation multiplies grid electricity consumption in MWh by the grid emission factor for the year and region.",
        "Scope 1 covers direct combustion the company controls, such as diesel burned in its own generator or forklift; the grid emission factor does not apply to it.",
    ],
    "reference-tnb-tariff-rp4.txt": [
        "PENINSULAR MALAYSIA ELECTRICITY TARIFF - REAL PUBLISHED REFERENCE DATA",
        "Figures below are real. Source: Tenaga Nasional Berhad and Energy Commission announcements on Regulatory Period 4, as reported in Malaysian press releases retrieved 2026-08-23. Content was rephrased for compliance with licensing restrictions.",
        "Regulatory Period 4 (RP4) runs from 2025 to 2027.",
        "The RP4 base tariff for Peninsular Malaysia was set at 45.40 sen/kWh.",
        "That figure was a revision of the 45.62 sen/kWh approved in December 2024.",
        "The preceding base tariff was 39.95 sen/kWh, so RP4 represents an increase.",
        "A tariff is a price per unit of electricity. It is not an emission factor and cannot be substituted for one.",
    ],
    "reference-sedg-disclosure-framework.txt": [
        "MALAYSIAN ESG DISCLOSURE FRAMEWORKS - REAL PUBLISHED REFERENCE DATA",
        "Sources: Capital Markets Malaysia, Simplified ESG Disclosure Guide (SEDG) Version 2, July 2025; Bursa Malaysia Sustainability Reporting Guide. Retrieved 2026-08-24. Summarised, not reproduced.",
        "SEDG is issued by Capital Markets Malaysia, an affiliate of the Securities Commission Malaysia, for SMEs answering ESG data requests from customers in their supply chain.",
        "SEDG Version 2 contains 38 disclosures across 15 topics and three pillars, and supersedes Version 1 which had 35.",
        "Version 2 added three disclosures: SEDG-E1.7, SEDG-E3.1 and SEDG-S2.3.",
        "Disclosures are graded Basic, Intermediate and Advanced so a company can start small and extend later.",
        "SEDG is voluntary, and it is aligned to the Bursa Malaysia Sustainability Reporting Guide, GRI, the ISSB standards (IFRS S1 and S2) and the GHG Protocol.",
        "Version 2 is aligned with the ASEAN Simplified ESG Disclosure Guide launched by the ASEAN Capital Markets Forum in April 2025.",
        "Bursa Malaysia requires Main Market listed issuers to report against nine common sustainability matters: anti-corruption, community and society, diversity, energy management, health and safety, labour practices and standards, supply chain management, data privacy and security, and water.",
        "ACE Market issuers report eleven common sustainability matters: the nine above plus waste management and emissions management.",
        "A listed customer that must report supply chain management is the reason an SME supplier receives a questionnaire like this one at all.",
    ],
    # ---- A tier: sound ----------------------------------------------------
    "A-06-safety-performance-fy2025.txt": [
        f"OCCUPATIONAL SAFETY AND HEALTH PERFORMANCE - {COMPANY} - SYNTHETIC SAMPLE",
        f"Reporting period: {PERIOD}. Scope: Klang plant only.",
        "Work-related fatalities in the reporting period: 0.",
        "Work-related injuries recorded in the reporting period: 4.",
        "Of those 4 injuries, 3 were lost-time injuries and 1 was a medical treatment case with no lost time.",
        "Total hours worked by all employees and contractors in the reporting period: 561,000.",
        "Lost time injury frequency rate (LTIFR): 5.35 per million hours worked, calculated as 3 lost-time injuries divided by 561,000 hours, multiplied by 1,000,000.",
        "The LTIFR denominator is stated above so the rate can be recomputed rather than taken on trust.",
        "Employees trained on health and safety standards in the reporting period: 268 of 268, which is 100 per cent.",
        "The company holds ISO 45001:2018 certification for the Klang plant, valid to 30 September 2027.",
        "Data owner: Safety and Health Officer, verified against the DOSH accident register on 12 January 2026.",
    ],
    "A-07-workforce-and-diversity-fy2025.txt": [
        f"WORKFORCE PROFILE AND DIVERSITY - {COMPANY} - SYNTHETIC SAMPLE",
        f"Reporting period: {PERIOD}. Scope: Klang plant only. All figures as at 31 December 2025.",
        f"Total employees: {HEADCOUNT}.",
        "Gender split of total employees: 96 women (35.8 per cent) and 172 men (64.2 per cent).",
        "Age band under 30: 74 employees (27.6 per cent).",
        "Age band 30 to 50: 158 employees (59.0 per cent).",
        "Age band over 50: 36 employees (13.4 per cent).",
        "Management positions: 24 in total, of which 9 are held by women (37.5 per cent).",
        "Employment type: 241 permanent, 27 fixed-term contract. No employees are engaged through a labour agent.",
        "Employees who left during the reporting period: 34.",
        "Employee turnover rate for the reporting period: 12.7 per cent, calculated as 34 leavers divided by 268 employees at period end.",
        "Child labour incidents recorded in the reporting period: 0. The company does not employ any person under 18.",
        "Forced labour incidents recorded in the reporting period: 0. No employee pays a recruitment fee and no employee's passport is held by the company.",
        "Data owner: Human Resources Manager. Reconciled against the December 2025 payroll run.",
    ],
    # ---- B tier: uncertain -------------------------------------------------
    "B-01-water-bills-q1-q2-fy2025.txt": [
        f"WATER WITHDRAWAL - PARTIAL RECORD - {COMPANY} - SYNTHETIC SAMPLE",
        "INCOMPLETE: this file covers January to June 2025 only. Statements for July to December 2025 have not been retrieved from the utility portal.",
        "Scope: Klang plant only. Source: purchased municipal supply (Air Selangor).",
        "January 2025: 682 cubic metres.",
        "February 2025: 641 cubic metres.",
        "March 2025: 714 cubic metres.",
        "April 2025: 699 cubic metres.",
        "May 2025: 706 cubic metres.",
        "June 2025: 678 cubic metres.",
        "Total for the six months shown: 4,120 cubic metres.",
        "No breakdown by source is available beyond purchased municipal supply; the site has no groundwater abstraction and no surface water intake.",
        "A full-year withdrawal figure cannot be derived from this file. Doubling the half-year total would be an estimate, not a measurement, and the second half of the year includes the annual shutdown in August.",
    ],
    "B-02-group-energy-summary-fy2025.txt": [
        "GROUP ENERGY SUMMARY - SCOPE DOES NOT MATCH THIS QUESTIONNAIRE - SYNTHETIC SAMPLE",
        "Prepared by: Tenggara Holdings Sdn. Bhd. group finance, February 2026.",
        "SCOPE WARNING: these figures cover all three group sites - the Klang plant, the Ipoh plant and the Port Klang logistics depot - combined. The questionnaire asks about the Klang plant only.",
        "Group total electricity purchased in FY2025: 4,912,600 kWh.",
        "Group total diesel consumed in FY2025: 41,800 litres.",
        "Group total LPG consumed in FY2025: 19,300 kilograms.",
        "No site-level breakdown is given in this summary.",
        "The group figure cannot be attributed to the Klang plant without a site breakdown, and the Klang plant's own metered figure is in A-01.",
    ],
    "B-03-recycled-content-estimate-fy2025.txt": [
        "RECYCLED INPUT MATERIAL - INTERNAL ESTIMATE, NOT VERIFIED - SYNTHETIC SAMPLE",
        "Prepared by: Purchasing executive, working note, 8 January 2026. Not reviewed or approved.",
        "ESTIMATE ONLY: the percentages below are the purchasing team's working assumption, not a supplier-confirmed figure.",
        "Aluminium extrusion: assumed 30 to 40 per cent recycled content, based on a verbal statement from one supplier.",
        "Steel fasteners: recycled content unknown. The supplier has not responded to two requests.",
        "Polymer housings: assumed 0 per cent recycled content.",
        "Overall recycled input material across all purchased materials: estimated at roughly one third, expressed as a range rather than a figure.",
        "No supplier has provided a certificate or a mass-balance record supporting any of these numbers.",
        "This note should not be quoted to a customer as a disclosed percentage until at least the aluminium and steel figures are confirmed in writing.",
    ],
    "B-04-energy-efficiency-initiatives-fy2025.txt": [
        f"ENERGY EFFICIENCY INITIATIVES - NO UNIT ON THE SAVING - {COMPANY} - SYNTHETIC SAMPLE",
        f"Reporting period: {PERIOD}. Scope: Klang plant.",
        "Initiative 1: LED replacement across the production hall, completed March 2025.",
        "Initiative 2: compressed air leak survey and repair programme, completed July 2025.",
        "Initiative 3: chiller scheduling changed to avoid peak-period operation, from September 2025.",
        "Combined saving reported by the maintenance team: 118,000.",
        "AMBIGUITY: the saving above is recorded without a unit. The maintenance log does not say whether 118,000 is kWh saved or ringgit saved, and at the RP4 tariff the two readings differ by roughly a factor of two.",
        "No baseline year is stated, so it is also unclear what the saving is measured against.",
        "The question asks for a reduction in kWh. This file cannot answer it as written.",
    ],
    # ---- C tier: wrong -----------------------------------------------------
    "C-01-weighbridge-waste-summary-fy2025.txt": [
        f"WEIGHBRIDGE WASTE SUMMARY - CONTRADICTS A-03 - {COMPANY} - SYNTHETIC SAMPLE",
        f"Reporting period: {PERIOD}. Scope: Klang plant. Source: gatehouse weighbridge tickets.",
        "Total scheduled waste despatched in FY2025: 18.4 metric tonnes.",
        "Total waste generated in FY2025: 229.5 metric tonnes.",
        "Waste diverted from disposal in FY2025: 131.0 metric tonnes.",
        "Waste directed to disposal in FY2025: 98.5 metric tonnes.",
        "CONTRADICTION: A-03, the DOE consignment-note summary, reports 12.6 tonnes of scheduled waste and 214.7 tonnes generated for the same site and the same period. This file reports 18.4 and 229.5.",
        "Both files are internally plausible and neither is obviously the more credible one. The weighbridge measures what leaves the gate; the consignment notes record what was formally consigned. They should reconcile and they do not.",
        "Nothing in the system will flag this. The rule engine only reports CONFLICTING when two evidence links carry different values for the same scope and period, and the keyword matcher never extracts a value at all. A reviewer has to notice.",
    ],
    "C-02-energy-and-emissions-fy2022.txt": [
        f"ENERGY AND EMISSIONS - WRONG YEAR - {COMPANY} - SYNTHETIC SAMPLE",
        "Reporting period: FY2022 (1 January 2022 to 31 December 2022). THIS IS NOT THE PERIOD THE QUESTIONNAIRE ASKS ABOUT.",
        "Total electricity purchased in FY2022: 1,612,400 kWh.",
        "Scope 2 emissions FY2022: 1,247.0 tCO2e, using the 0.774 Peninsular factor for 2022.",
        "Scope 1 emissions FY2022: 51.3 tCO2e.",
        "Total waste generated FY2022: 188.2 metric tonnes.",
        "Average training hours per employee FY2022: 9.4.",
        "STALE: these figures are three years older than the FY2025 period the questionnaire covers. They are presented in the same layout as the current-year files and are easy to cite by mistake.",
        "The system will not mark this outdated either: documents.source_date cannot be set on upload, so the rule engine has no date to test against the 24-month threshold.",
    ],
    "C-03-safety-record-wrong-entity.txt": [
        "OCCUPATIONAL SAFETY RECORD - WRONG LEGAL ENTITY - SYNTHETIC SAMPLE",
        "Entity: Tenggara Logistics Sdn. Bhd. THIS IS A DIFFERENT COMPANY FROM THE ONE THE QUESTIONNAIRE IS ABOUT.",
        "Tenggara Logistics Sdn. Bhd. is a separate legal entity under the same holding company, operating the Port Klang depot.",
        f"Reporting period: {PERIOD}. Scope: Port Klang depot.",
        "Work-related fatalities: 1. A fatal forklift incident occurred at the depot on 6 August 2025.",
        "Work-related injuries: 11.",
        "Lost time injury frequency rate (LTIFR): 14.2 per million hours worked.",
        "WRONG SCOPE: citing this file against the questionnaire would report a fatality that did not happen at the site being asked about, and would roughly treble the injury count.",
        "The keyword matcher will happily match this file to the safety question. Nothing in the pipeline compares the entity name in the document against the case.",
    ],
    "C-04-data-privacy-note-contradicts-itself.txt": [
        f"CUSTOMER DATA AND PRIVACY - INTERNALLY INCONSISTENT - {COMPANY} - SYNTHETIC SAMPLE",
        f"Reporting period: {PERIOD}. Scope: Klang plant.",
        "Substantiated complaints concerning breaches of customer privacy in FY2025: none.",
        "Losses of customer data in FY2025: none recorded.",
        "Incident log extract, 19 September 2025: a laptop containing a customer contact list was reported stolen from a vehicle. The customer was notified on 22 September 2025.",
        "Incident log extract, 4 November 2025: the September laptop incident was closed after the device was confirmed encrypted.",
        "CONTRADICTION WITHIN ONE FILE: the summary lines say there were no losses of customer data, and the incident log below them describes one that was notified to the customer.",
        "Whether the September incident is a 'loss of customer data' for disclosure purposes is a judgement call, and it is precisely the judgement a reviewer is supposed to make rather than a matcher.",
    ],
}


def build_text_files() -> None:
    for name, lines in TEXT_FILES.items():
        (EVIDENCE_DIR / name).write_text("\n".join(lines) + "\n", encoding="utf-8")


# --------------------------------------------------------------------------- #


def main() -> None:
    QUESTIONNAIRE_DIR.mkdir(parents=True, exist_ok=True)
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)

    build_questionnaire()
    build_broken_questionnaire()
    build_electricity_bills_pdf()
    build_unreadable_pdf()
    build_ghg_inventory_xlsx()
    build_scheduled_waste_xlsx()
    build_training_register_xlsx()
    build_anti_bribery_policy_docx()
    build_superseded_handbook_docx()
    build_text_files()

    for directory in (QUESTIONNAIRE_DIR, EVIDENCE_DIR):
        rel = directory.relative_to(ROOT)
        for path in sorted(directory.iterdir()):
            print(f"  {rel}/{path.name}  ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
